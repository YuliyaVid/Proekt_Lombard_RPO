import sys  # sys нужен для передачи argv в QApplication
import os  as os# Отсюда нам понадобятся методы для отображения содержимого директорий
import sqlite3 # библиотека для работы с Базой данных
import docx # библиотека для работы с doc файлами

from PyQt5       import QtWidgets, QtGui, QtCore
from PyQt5.QtGui import QBrush, QColor  #импорт библиотеки для работы с оконным интерфейсом
import datetime # библиотека для получения текущей даты
import design_login_window # дизайн окна ввода пароля
import design_worker_window # дизайн рабочего окна сотрудника
import design_addc_window # дизайн окна добавления клиентов
import design_addc_2_window # дизайн окна изменения данных клиентов
import design_addim_window # дизайн окна добавления имущества
import design_doc_window # дизайн окна готового договора



class Login_app(QtWidgets.QMainWindow, design_login_window.Ui_A): #класс окна ввода пароля
    def __init__(self): 
        self.conn = sqlite3.connect(r'db/db.db') # подключение к базе данных которая находится в папке db
        self.cur = self.conn.cursor() # создание курсора базы данных
        self.count_of_loggins = 0 # переменная количества попыток входа (для ограничения попыток входа

        super().__init__() # 
        self.setupUi(self)  # эта и предыдущая строка необходимы для отображения нашего окна
        self.Registration_button.clicked.connect(self.reg) # событие нажатия кнопки регистрация  (см ниже описание функции def reg(self))
        self.Login_button.clicked.connect(self.login)  # событие нажатия кнопки авторизация (см ниже описание функции def login(self))



    def reg(self): # функция которая срабатывает при нажатии на кнопку регистрация
        self.cur.execute('SELECT id FROM workers') # для того чтобы зарегистрировать пользователя ему необходимо задать ID . для этого мы выбираем все ID которые у нас есть
        rows = self.cur.fetchall() # сохранияем все id в переменной rows 
        self.usr=(len(rows)+1, self.lineEdit.text(), self.lineEdit_2.text()) # формируем кортеж данных для загрузки в БД где первый элемент наш НОВЫЙ id посчитанный пряму здесь (длина rows == кол во id) + 1
        self.cur.execute(f"INSERT INTO workers VALUES(? ,?, ?);", self.usr) # загружаем наш кортеж данных второй и третий элемент кортежа это логин и пароль взятых с полей lineEdit и LineEdit_2
        self.conn.commit() # подтверждаем изменения в БД

    def login(self): # функция которая срабатывает при нажатии на кнопку авторизация
        self.usr=(self.lineEdit.text(), self.lineEdit_2.text()) # формируем кортеж данных логин + пароль (информацию берем с тех же полей ввода что и в предыдущей функции)
        password = self.lineEdit_2.text() 
        self.cur.execute("SELECT COUNT(*) from workers where name=? AND password = ?", self.usr) # ищем количество сопадений Логин + Пароль в таблице workers (в нашей базе данных). 
        result=self.cur.fetchone() # результат поиска записываем в result 
        if result[0] == 1: # если совпадение есть то 
            self.next = Worker_app(self.lineEdit.text()) # создаем экземпляр класса окна Worker_app (см ниже) а также передаем ему имя пользователя
            self.next.show() # запускаем окно воркер ап
            self.conn.close() # закрываем соединение с БД
            self.close() #закрываем окно ввода пароля
            return(True)

        else: # если логин и (или) пароль введены не верно то
            self.count_of_loggins += 1  # увеличиваем количество попыток ввода
            self.label_3.setText(f"Попытка {self.count_of_loggins} из 3") # пишем пользователю о том сколько попыток он сделал и сколько у него осталось
            if self.count_of_loggins == 3: # если количество попыток 3, то
                raise SystemExit(1) # закрываем нашу программу

class Worker_app(QtWidgets.QMainWindow, design_worker_window.Ui_Main_window): # основное окно программы
    def __init__(self,name):
        super().__init__()
        self.price_list = [] # переменная которая получаем прайс лист
        self.name_worker = name # сохраняем имя пользователя с предыдущего окна
        self.setupUi(self) # запускаем окно
        self.worker_place.setTitle('Сотрудник - ' + self.name_worker) # в рамочке где сотрудник пишем имя сотрудника 
        
        self.conn = sqlite3.connect(r'db/db.db') # подключаемся к базе данных
        self.cur = self.conn.cursor() # показываем курсор

        rows = self.cur.execute('SELECT * FROM clients').fetchall() # выбираем все поля с таблицы клиенты 
        for row in rows: # читаем каждую строчку из таблицы клиенты
            self.listWidget.addItem(f'{str(row[0])}  {str(row[1])}  {str(row[2])}') # в списке который в окне справа добавляем клиентов из базы данных. 0 это айди 1 это фамилия 2 это имя

        rows = self.cur.execute('SELECT * FROM imushestvo').fetchall() # выбираем все поля с таблицы имущество
        for row in rows: # читаем каждую строчку из таблицы имущество
            self.listWidget_2.addItem(f'{str(row[0])}  {str(row[2])}  {str(row[3])}') # в списке который в окне справа добавляем имущество из базы данных

        self.listWidget.clicked[QtCore.QModelIndex].connect(self.on_clicked) # если мы нажимаем на конкретного клиента создается событие которое обрабатывается функцией on_clicked
        self.Add_item.clicked.connect(self.Add_item_click) # событие нажатия кнопки добавления имещуства обрабатывается функцией add item 
        self.ref_but.clicked.connect(self.ref_but_click) # событие нажатия кнопки обновления списка клиентов обрабатывается функцией ref_but_click
        self.Add_client.clicked.connect(self.Add_client_click) #событие нажатия кнопки добавления клиентов обрабатывается функцией add_client_click
        self.Dell_client.clicked.connect(self.dell_c) # событие нажатия кнопки удаления клиента (см ниже описание функции def dell_c))
        self.Add_client_2.clicked.connect(self.add_c_2) # событие нажатия кнопки изменения данных о клиенте
        self.Exit_button.clicked.connect(self.exit_app) # событие нажатия кнопки выход

        f = open('price_list.txt', encoding = 'utf-8') # открываем файл прайс лист тхт
        for line in f.readlines(): #считываем строки
            line = line.split() #разделяем строки по пробелу к примеру была строка "мебель 1000" станет списком ["мебель","1000"]
            self.price_list.append(line) #добавляем в список прайс листа наши строки-списки

    def exit_app(self): # функция выхода из программы
        raise SystemExit(1)# закрываем нашу программу

    def ref_but_click(self): # функция обновления списка клиентов (необходимо при добавлении нового клиента чтобы он отобразился в списке)
        self.listWidget.clear() # очищаем список клиентов
        self.listWidget_2.clear() # очищаем список имущества
        rows = self.cur.execute('SELECT * FROM clients').fetchall() # выбираем всех клиентов в базе данных 
        for row in rows: 
            self.listWidget.addItem(f'{str(row[0])}  {str(row[1])}  {str(row[2])}') # вставляем в список всех клиентов как и до этого

        rows = self.cur.execute('SELECT * FROM imushestvo').fetchall() # выбираем все поля с таблицы имущество 
        for row in rows: # читаем каждую строчку из таблицы имущество
            self.listWidget_2.addItem(f'{str(row[0])}  {str(row[2])}  {str(row[3])}') # в списке который в окне справа добавляем имущество из базы данных. 


        self.Add_item.setEnabled(False) # так как мы обновили списки теперь нам необходимо заново выбрать клиента (и погасить кнопки)
        self.Dell_client.setEnabled(False) # так как мы обновили списки теперь нам необходимо заново выбрать клиента (и погасить кнопки)
        self.Add_client_2.setEnabled(False)  # так как мы обновили списки теперь нам необходимо заново выбрать клиента (и погасить кнопки)


    def on_clicked(self): # если мы выбираем клиента в списке который справа
        self.Add_item.setEnabled(True) # то кнопка добавить имущество становится активной
        self.Dell_client.setEnabled(True) # нопка удаления клиента становится активной
        self.Add_client_2.setEnabled(True)  # кнопка изменения данных о клиенте становится активной

    def Add_client_click(self): # функция добавления клиента вызываемая по кнопке добавить клиента
        self.next1 = Client_add() # создаем экземпляр класса Client_add (окно добавления клиента)
        self.next1.show() # запускаем окно добавления клиента

    def add_c_2(self): # функция которая вызывает окно для изменения данных о клиенте
        self.client_info_in_list = self.listWidget.currentItem().text().split() # берем строку из списка клиентов и разбиваем ее например была строка "2 семенов алексей" стала "[2, 'семенов', 'алексей']" 
        self.cur.execute("SELECT * from clients where id= ?", self.client_info_in_list[0]) #выбираем полную информацию о клиенте по id
        self.full_info_of_c = self.cur.fetchone() # сохраняем всю информацию о клиенте в переменной full_indo_of_c
        self.next = Add_c_2(self.full_info_of_c) # создаем окно Add_c_2 изменения данных о клиенте
        self.next.show() # показываем окно изменения данных о клиенте

        

    def Add_item_click(self): # функция добавления имущества 
        self.client_info_in_list = self.listWidget.currentItem().text().split() # берем строку из списка клиентов и разбиваем ее например была строка "2 семенов алексей" стала "[2, 'семенов', 'алексей']" 
        self.cur.execute("SELECT * from clients where id= ?", self.client_info_in_list[0]) #выбираем полную информацию о клиенте по id
        self.full_info_of_c = self.cur.fetchone() # сохраняем всю информацию о клиенте в переменной full_indo_of_c
        self.next = Im_add(self.price_list, self.listWidget.currentItem().text().split(), self.full_info_of_c) # создаем окно Im_add (добавление имущества)
        self.next.show() # показываем окно добавления имущества

    def dell_c(self): #функция удаление клиента
        self.client_info_in_list = self.listWidget.currentItem().text().split() # берем строку из списка клиентов и разбиваем ее например была строка "2 семенов алексей" стала "[2, 'семенов', 'алексей']" (в целом тут нам нужен только ID)
        self.cur.execute("DELETE FROM clients WHERE id = ?", self.client_info_in_list[0]) # удаляем запись о клиенте (пользуемся тем что id у нас уникальный)
        self.conn.commit() # подтверждаем изменения в базе данных
        self.ref_but_click() # пользуемся готовой фунцией которую мы реализовали для обновления списков чтобы получить актуальную информацию



class Client_add(QtWidgets.QMainWindow, design_addc_window.Ui_MainWindow): #окно добавления клиента
    def __init__(self): 
        super().__init__()
        self.setupUi(self)

        self.client_list = [] # список состоящий из данных о клиенте (пока пустой)

        self.conn = sqlite3.connect(r'db/db.db') # подключение к базе данных
        self.cur = self.conn.cursor() # создание курсора

        self.cur.execute('SELECT id FROM clients') # выборка из базы данных по полю id таблицы clients 
        rows = self.cur.fetchall() # сохранения данных в переменной rows 
        self.cl_id = len(rows) + 1 # создание нового id клиента путем прибавления к кол-ву айдишников единицы
        self.lineEdit.setText(str(self.cl_id)) # указание нового айди в первом поле 
        self.Add_client_f.clicked.connect(self.Add_client_click) # событие нажатия на кнопку "добавить клиента" вызывает функцию Add_client_click

    def Add_client_click(self): # событие добавления пользователя в базу данных
        self.client_list = [self.cl_id, self.lineEdit_2.text(), self.lineEdit_3.text(),
            self.lineEdit_4.text(), self.lineEdit_5.text(), self.lineEdit_6.text(),
            self.lineEdit_7.text(), self.lineEdit_8.text(), self.lineEdit_9.text()
            ] # формирование записи вида [id, surname, name, otchestvo, birthday, pasS, pasN, adress, telnumber ]
        self.cur.execute(f"INSERT INTO clients VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?);", self.client_list) # добавляем сформированную запись в базу данных
        self.conn.commit() # подтверждаем изменения в базе данных
        self.close() # закрываем окно

class Add_c_2(QtWidgets.QMainWindow, design_addc_2_window.Ui_MainWindow): #окно изменения данных о клиенте
    def __init__(self, full_info_of_c):
        super().__init__()
        self.setupUi(self) # отображаем интерфейс

        self.conn = sqlite3.connect(r'db/db.db') # подключение к базе данных
        self.cur = self.conn.cursor() # создание курсора

        self.client_list = [] # список состоящий из данных о клиенте (пока пустой)

        self.full_info_of_c = full_info_of_c                 
        self.lineEdit.setText(str(self.full_info_of_c[0]))
        self.lineEdit_2.setText(str(self.full_info_of_c[1]))  # загружаем все данные с нашей базы данных в поля для редактирования
        self.lineEdit_3.setText(str(self.full_info_of_c[2]))
        self.lineEdit_4.setText(str(self.full_info_of_c[3]))
        self.lineEdit_5.setText(str(self.full_info_of_c[4]))
        self.lineEdit_6.setText(str(self.full_info_of_c[5]))
        self.lineEdit_7.setText(str(self.full_info_of_c[6]))
        self.lineEdit_8.setText(str(self.full_info_of_c[7]))
        self.lineEdit_9.setText(str(self.full_info_of_c[8]))

        self.Add_client_f.clicked.connect(self.Add_client_click) # событие нажатия на кнопку "Изменить данные" вызывает функцию Add_client_click
    def Add_client_click(self): # функция изменения данных
        self.client_list = [self.lineEdit_2.text(), self.lineEdit_3.text(),
            self.lineEdit_4.text(), self.lineEdit_5.text(), self.lineEdit_6.text(),
            self.lineEdit_7.text(), self.lineEdit_8.text(), self.lineEdit_9.text()
            ] # формирование записи вида [surname, name, othername, birsthday, spas, npas, adress, telnumber ]
        self.cur.execute(f"UPDATE clients SET surname = ?, name = ?, othername = ?, birsthday = ?, spas = ?, npas = ?, adress = ?, telnumber = ? WHERE id = {self.full_info_of_c[0]};", self.client_list) # добавляем сформированную запись в базу данных
        self.conn.commit() # подтверждаем изменения в базе данных
        self.close() # закрываем окно

        
        
class Im_add(QtWidgets.QMainWindow, design_addim_window.Ui_MainWindow): # окно добавления имущества
    def __init__(self, price, client_info, full_info_of_c):
        super().__init__()
        self.setupUi(self)
        self.full_info_of_c = full_info_of_c # сохранения данных о клиенте 
        self.price = price # сохранения прайс листа
        self.client_info = client_info # тоже данные о клиенте
        self.items_list = ['Выберите тип'] # добавление строчки "выберите тип" в выпадающий список
        self.im_id = 0 # создание переменной айди имущества
        for i in range(len(self.price)): #
            self.items_list.append(self.price[i][0]) # формирование списка состоящего только из названий прайс листа

        self.conn = sqlite3.connect(r'db/db.db') # подключение к базе данных
        self.cur = self.conn.cursor() # создание курсора базы данных
        self.refresh() # обновление страницы (см функцию refresh ниже)

        self.comboBox.activated.connect(self.handleActivated) # событие выбора элемента из выпадающего списка
        self.pushButton.clicked.connect(self.Add_item_click) # событие нажатия на кнопку добавить имущество
        

    def handleActivated(self, index):
        self.type = self.comboBox.itemText(index) # при выборе элемента из списка мы в self.type засовываем значение текста выпадающего списка

    def Add_item_click(self):
        self.item = (self.im_id, self.client_info[0], self.type, self.lineEdit_4.text(), self.lineEdit_5.text()) # формирование записи о добавлении имущества
        self.cur.execute(f"INSERT INTO imushestvo VALUES(?, ?, ?, ?, ?);", self.item) # добавление записи об имуществе в табоицу "имущество"
        self.conn.commit() # подтверждение изменений
        

        self.next = doc_show(self.price, self.full_info_of_c, self.lineEdit_6.text(), self.item) # отображение сформированного договора
        self.next.show() # показ окна с договором
 
        self.refresh() # обновление страницы
        self.conn.close() # закрытие базы
        self.close() # закрытие окна

    def refresh(self): # обновление страницы 
        self.comboBox.clear() # очищение списка
        self.comboBox.addItems(self.items_list) # добаление элементов в список 
        self.lineEdit_2.setText(self.client_info[0]) # вывод id клиента
        self.cur.execute('SELECT id FROM imushestvo') # сбор количества айди в таблице имущество
        rows = self.cur.fetchall() # продолжение строки выше + сохранение айдишников в переменной ровс
        self.im_id = len(rows) + 1 #  создание нового айди имущества
        self.lineEdit.setText(str(self.im_id)) # вывод айди имущества в первой строке
        self.lineEdit_4.setText('') # очистка полей ввода
        self.lineEdit_5.setText('')  #
        self.lineEdit_6.setText('')   #

class doc_show(QtWidgets.QMainWindow, design_doc_window.Ui_MainWindow): # окно формирования договора 
    def __init__(self, price, full_info_of_c, date, item):
        super().__init__()
        self.setupUi(self)

        self.client = full_info_of_c # получение данных о клиенте
        self.date = date # получение даты сроков хранения имущества
        self.item = item # получение данных об имуществе 

        for i in range(len(price)):
            if price[i][0] == self.item[2]:
                self.price = price[i][1] # получение данных о стоимости данного типа имущества

        now = datetime.datetime.now() # получение текущей даты
        self.conn = sqlite3.connect(r'db/db.db') # подключение к базе данных 
        self.cur = self.conn.cursor() # подключение курсора
        self.cur.execute('SELECT id FROM doc') # выборка из базы данных по айди документов
        rows = self.cur.fetchall() # сохранение айдишников в ровс
        self.doc_id = len(rows) + 1 # создание нового айдишника 


        doc = docx.Document() # создание объекта докх документ (все что ниже это формирование строк документа)
        par1 = doc.add_paragraph('              Договор\n')
        par1.add_run('          хранения вещей в ломбарде\n')
        par1.add_run(f'№____{self.doc_id}_________                   "{now.day}"___{now.month}____ {now.year}г.\n')
        par1.add_run('"Ломбард", в лице _____________________________________________________,\n')
        par1.add_run(f' с одной стороны, и _____________{self.client[1]}_{self.client[2]}_{self.client[3]}______________________,\n')
        par1.add_run('                                     (Ф.И.О. гражданина)\n')
        par1.add_run('    именуемый в   дальнейшем   "Клиент",   с  другой  стороны,  заключили\n')
        par1.add_run('    настоящий договор о нижеследующем.\n')
        par1.add_run('         1. По  настоящему  договору  Ломбард  обязуется  хранить   вещь,\n')
        par1.add_run('    переданную ему Клиентом, и возвратить эту вещь в сохранности.\n')
        par1.add_run(f'         2. На хранение в Ломбард сдается: ____________{self.item[2]}___________\n')
        par1.add_run(f'    _{self.item[3]}_____{self.item[4]}__________________________________.\n')
        par1.add_run('            (наименование вещи и ее индивидуализирующие признаки)\n')
        par1.add_run('         3. Вещь,   сдаваемая   на   хранение,   оценена   Сторонами    в\n')
        par1.add_run(f'    установленном порядке в __________{self.price}______________________________.\n')
        par1.add_run('                                         (сумма)\n')
        par1.add_run(f'         3. Настоящий договор заключен сроком до: _______{self.date}__________.\n')
        par1.add_run('         4. За  хранение  вещи  взимается  плата  в следующем размере \n')
        par1.add_run(f'    : ________________________________{float(self.price) - (float(self.price)*0.9)}________________________\n')
        par1.add_run('         5. Адреса и реквизиты сторон:\n')
        par1.add_run('         Ломбард: __________________OOO_Ломбард_________________________\n')
        par1.add_run(f'         Клиент: _{self.client[1]}_{self.client[2]}_{self.client[3]}_Дата_рождения____{self.client[4]}__________\n')
        par1.add_run('                      (фамилия, имя, отчество, дата рождения, телефон)\n')
        par1.add_run(f'    паспорт: серия ____{self.client[5]}____, Nо. ___{self.client[6]}_\n')
        par1.add_run(f'    адрес: ________________{self.client[7]}__________________\n')
        par1.add_run('                               Подписи сторон:\n')
        par1.add_run('             Ломбард                                     Клиент\n')
        par1.add_run('    ________________________                     ________________________\n')
        par1.add_run('              \n')
        doc.save(f'doc_id_{self.doc_id}.docx') # сохранение документа под индивидуальным айди

        self.doc_list = [self.doc_id, self.item[0], f'{now.year}-{now.month}-{now.day}', self.date, self.price, 0.9] # создание записи о документе для добавления его в базу
        self.cur.execute(f"INSERT INTO doc VALUES(?, ?, ?, ?, ?, ?);", self.doc_list) # добавление документа в базу
        self.conn.commit() # подтверждение изменений в базе

        doc = docx.Document(f'doc_id_{self.doc_id}.docx') # открываем документ
        for i in doc.paragraphs:
            self.textEdit.insertPlainText(i.text) # читаем построчно и выводим это в наше поле

        self.pushButton.clicked.connect(self.ok_click) # если нажимаем на кнопку ОК

    def ok_click(self): # то выходим
        self.close() # здесь




       
    
        

def main():
    app = QtWidgets.QApplication(sys.argv)  # Новый экземпляр QApplication
    window = Login_app()  # Создаём объект класса ExampleApp
    window.show()  # Показываем окно
    app.exec_()  # и запускаем приложение

        

if __name__ == '__main__':  # Если мы запускаем файл напрямую, а не импортируем
    main()  # то запускаем функцию main()